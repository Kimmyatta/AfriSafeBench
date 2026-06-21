import json
import re
from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
RESULTS_DIR = ROOT / "data" / "afrisafebench" / "results"
FIGURES_DIR = ROOT / "docs" / "figures"
REPORTS_DIR = ROOT / "docs" / "reports"
SUBMISSION_PATH = ROOT / "SUBMISSION.md"

PALETTE = {
    "teal": (31, 118, 109),
    "gold": (213, 159, 42),
    "red": (201, 79, 79),
    "blue": (61, 111, 182),
    "ink": (23, 33, 43),
    "muted": (96, 112, 134),
    "grid": (216, 224, 234),
    "light": (237, 242, 246),
    "paper": (255, 255, 255),
}


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def draw_text(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, size: int, *, bold: bool = False, fill=None) -> None:
    draw.text(xy, text, font=font(size, bold), fill=fill or PALETTE["ink"])


def draw_wrapped(draw: ImageDraw.ImageDraw, x: int, y: int, text: str, width: int, size: int = 24, fill=None) -> int:
    line_height = int(size * 1.25)
    for line in wrap(text, width=width):
        draw_text(draw, (x, y), line, size, fill=fill or PALETTE["muted"])
        y += line_height
    return y


def save_model_coverage(summary: dict) -> Path:
    values = summary["mean_coverage_score_by_model"]
    order = sorted(values, key=values.get, reverse=True)
    image = Image.new("RGB", (1400, 640), PALETTE["paper"])
    draw = ImageDraw.Draw(image)
    draw_text(draw, (60, 55), "Mean Coverage by Model", 42, bold=True)
    draw_text(draw, (60, 105), "AfriSafeBench benchmark: 25 scenarios x 3 models = 75 evaluations", 24, fill=PALETTE["muted"])

    colors = [PALETTE["teal"], PALETTE["blue"], PALETTE["gold"]]
    left, top, bar_w, bar_h, gap = 430, 190, 780, 68, 48
    for index, model in enumerate(order):
        y = top + index * (bar_h + gap)
        value = values[model]
        draw_text(draw, (60, y + 18), model, 27, bold=True)
        draw.rounded_rectangle((left, y, left + bar_w, y + bar_h), radius=8, fill=PALETTE["light"])
        draw.rounded_rectangle((left, y, left + int(bar_w * value / 100), y + bar_h), radius=8, fill=colors[index])
        draw_text(draw, (left + int(bar_w * value / 100) + 20, y + 18), f"{value:.2f}%", 27, bold=True)

    output = FIGURES_DIR / "model_coverage.png"
    image.save(output)
    return output


def save_category_weakness(summary: dict) -> Path:
    by_category = summary["coverage_score_by_risk_category"]
    categories = sorted({category for rows in by_category.values() for category in rows})
    averages = {
        category: sum(rows.get(category, 0) for rows in by_category.values()) / len(by_category)
        for category in categories
    }
    weakest = sorted(averages.items(), key=lambda item: item[1])[:10]

    image = Image.new("RGB", (1600, 1000), PALETTE["paper"])
    draw = ImageDraw.Draw(image)
    draw_text(draw, (60, 55), "Risk Categories with Lowest Average Coverage", 42, bold=True)
    draw_text(draw, (60, 105), "Average across evaluated models. Lower scores indicate harder-to-detect governance risks.", 24, fill=PALETTE["muted"])

    left, top, bar_w, bar_h, gap = 650, 180, 760, 46, 36
    for index, (category, value) in enumerate(weakest):
        y = top + index * (bar_h + gap)
        draw_wrapped(draw, 60, y + 3, category, width=36, size=23)
        draw.rounded_rectangle((left, y, left + bar_w, y + bar_h), radius=8, fill=PALETTE["light"])
        color = PALETTE["red"] if value < 50 else PALETTE["gold"] if value < 70 else PALETTE["teal"]
        draw.rounded_rectangle((left, y, left + int(bar_w * value / 100), y + bar_h), radius=8, fill=color)
        draw_text(draw, (left + bar_w + 20, y + 8), f"{value:.1f}%", 24, bold=True)

    output = FIGURES_DIR / "category_weakness.png"
    image.save(output)
    return output


def save_framework_priority(summary: dict) -> Path:
    values = summary.get("priority_distribution", {})
    image = Image.new("RGB", (1200, 560), PALETTE["paper"])
    draw = ImageDraw.Draw(image)
    draw_text(draw, (60, 55), "Framework Guidance Recommendation Priorities", 40, bold=True)
    draw_text(draw, (60, 105), "Representative sample: 5 scenarios x 3 models = 15 governance reports", 24, fill=PALETTE["muted"])

    max_count = max(values.values()) if values else 1
    colors = {"High": PALETTE["red"], "Medium": PALETTE["gold"], "Low": PALETTE["teal"]}
    left, top, bar_w, bar_h, gap = 320, 190, 650, 66, 46
    for index, priority in enumerate(["High", "Medium", "Low"]):
        count = values.get(priority, 0)
        y = top + index * (bar_h + gap)
        draw_text(draw, (60, y + 16), priority, 28, bold=True)
        draw.rounded_rectangle((left, y, left + bar_w, y + bar_h), radius=8, fill=PALETTE["light"])
        draw.rounded_rectangle((left, y, left + int(bar_w * count / max_count), y + bar_h), radius=8, fill=colors[priority])
        draw_text(draw, (left + bar_w + 20, y + 16), str(count), 28, bold=True)

    output = FIGURES_DIR / "framework_priority_distribution.png"
    image.save(output)
    return output


def save_workflow_diagram() -> Path:
    image = Image.new("RGB", (1800, 560), PALETTE["paper"])
    draw = ImageDraw.Draw(image)
    draw_text(draw, (60, 55), "AfriSafeBench Workflow", 42, bold=True)
    draw_text(draw, (60, 105), "Evaluation benchmark plus framework-guided governance tool.", 24, fill=PALETTE["muted"])

    steps = [
        ("Scenario", "African healthcare AI deployment case"),
        ("LLM Review", "Model identifies risks and severities"),
        ("Scoring", "Matched, missed, extra risks"),
        ("Human Audit", "Flags semantic rescoring changes"),
        ("Framework Retrieval", "WHO, NIST, UNESCO, OECD, AU chunks"),
        ("Recommendations", "Governance actions and checklist"),
    ]
    box_w, box_h, x0, y, gap = 245, 140, 60, 230, 42
    for index, (title, subtitle) in enumerate(steps):
        x = x0 + index * (box_w + gap)
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=14, fill=(247, 250, 252), outline=PALETTE["grid"], width=2)
        draw_text(draw, (x + 20, y + 24), title, 25, bold=True)
        draw_wrapped(draw, x + 20, y + 65, subtitle, width=23, size=20)
        if index < len(steps) - 1:
            ax = x + box_w + 10
            ay = y + box_h // 2
            draw.line((ax, ay, ax + 22, ay), fill=PALETTE["muted"], width=4)
            draw.polygon([(ax + 22, ay - 8), (ax + 38, ay), (ax + 22, ay + 8)], fill=PALETTE["muted"])

    output = FIGURES_DIR / "workflow_diagram.png"
    image.save(output)
    return output


def generate_figures() -> dict[str, Path]:
    FIGURES_DIR.mkdir(parents=True, exist_ok=True)
    benchmark = read_json(RESULTS_DIR / "benchmark_summary.json")
    framework = read_json(RESULTS_DIR / "framework_guidance_summary.json")
    return {
        "Mean coverage score by model:": save_model_coverage(benchmark),
        "Average coverage and miss rate by category across the three evaluated models:": save_category_weakness(benchmark),
        "Initial framework-guided reports produced recommendations such as:": save_framework_priority(framework),
        "The main contribution is a reusable workflow:": save_workflow_diagram(),
    }


def set_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"


def add_picture(document: Document, path: Path, caption: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.5))
    caption_paragraph = document.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run_item in caption_paragraph.runs:
        run_item.italic = True


def markdown_to_docx(markdown: str, output_path: Path, figure_map: dict[str, Path]) -> None:
    document = Document()
    set_styles(document)
    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    captions = {
        "Mean coverage score by model:": "Figure 1. Mean coverage by model.",
        "Average coverage and miss rate by category across the three evaluated models:": "Figure 2. Risk categories with lowest average coverage.",
        "Initial framework-guided reports produced recommendations such as:": "Figure 3. Framework guidance recommendation priorities.",
        "The main contribution is a reusable workflow:": "Figure 4. AfriSafeBench workflow.",
    }

    def flush_code() -> None:
        nonlocal code_lines
        if not code_lines:
            return
        paragraph = document.add_paragraph()
        run = paragraph.add_run("\n".join(code_lines))
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        code_lines = []

    def flush_table() -> None:
        nonlocal table_rows
        if not table_rows:
            return
        rows = table_rows
        if len(rows) >= 2 and all(set(cell.strip()) <= {"-", ":"} for cell in rows[1]):
            rows = [rows[0]] + rows[2:]
        table = document.add_table(rows=0, cols=len(rows[0]))
        table.style = "Table Grid"
        for row_index, row in enumerate(rows):
            cells = table.add_row().cells
            for column_index, value in enumerate(row):
                cells[column_index].text = value.strip()
                if row_index == 0:
                    for paragraph in cells[column_index].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
        table_rows = []

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.startswith("|") and line.endswith("|"):
            table_rows.append([cell.strip() for cell in line.strip("|").split("|")])
            continue
        if table_rows:
            flush_table()
        if not line.strip():
            document.add_paragraph("")
            continue
        if line.startswith("# "):
            paragraph = document.add_heading(line[2:].strip(), level=0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
            continue

        paragraph = document.add_paragraph()
        for part in re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", line):
            if not part:
                continue
            if part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1])
                run.font.name = "Courier New"
            elif part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

        if line in figure_map:
            add_picture(document, figure_map[line], captions[line])

    flush_code()
    flush_table()
    document.save(output_path)


def main() -> None:
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    figure_map = generate_figures()
    markdown = SUBMISSION_PATH.read_text(encoding="utf-8")
    output_path = REPORTS_DIR / "AfriSafeBench_Submission_Report_With_Figures.docx"
    markdown_to_docx(markdown, output_path, figure_map)
    print(output_path)


if __name__ == "__main__":
    main()
